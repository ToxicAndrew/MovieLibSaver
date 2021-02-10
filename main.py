import sys  # sys нужен для передачи argv в QApplication
from PyQt5 import QtWidgets  # Импортируем QtWidgets
from PyQt5.QtWidgets import *  # Импортируем все объекты QtWidgets

import mainwin  # Импортируем стили окон
import requests
from bs4 import BeautifulSoup
from docx import Document

HEADERS = {'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) '
                         'Chrome/86.0.4240.198 Safari/537.36 OPR/72.0.3815.487', 'accept': '*/*'}

class MainWindow(QtWidgets.QMainWindow, mainwin.Ui_MainWindow):
    def __init__(self, parent=None):
        super(MainWindow, self).__init__(parent)
        self.setupUi(self)
        self.go.clicked.connect(self.fgo)
        self.OpenPatch.clicked.connect(self.choose_path)
        conn = getHtml('https://www.movielib.ru').status_code
        if conn != 200:
            QMessageBox.critical(self, "Ошибка!", "Невозможно установить соединение с сайтом!")
            self.go.disconnect()


    def fgo(self):
        if self.Nickname.text() == '' and self.PatchToSave.text() == '':
            QMessageBox.warning(self, "Внимание!", "Заполните поля")
        elif self.Nickname.text() == '' and self.PatchToSave.text() != '':
            QMessageBox.warning(self, "Внимание!", "Введите свой ник!")
        elif self.Nickname.text() != '' and self.PatchToSave.text() == '':
            QMessageBox.warning(self, "Внимание!", "Выберите место сохранения файла!")
        else:
            main_adr = self.get_adress_hsf()
            check_nick = getHtml(main_adr).status_code
            if check_nick != 200:
                QMessageBox.critical(self, "Ошибка!", "Проверьте правильность ника!")
            else:
                hsm_html = main_adr + 'movies/'
                dcount = count(hsm_html)
                self.SeenFilmsVal.display(dcount)

                wlm_html = main_adr + 'wish/'
                dcount = count(wlm_html)
                self.WishFilmsVal.display(dcount)

                lm_html = main_adr + 'movies/love/'
                dcount = count(lm_html)
                self.LikeFilmsVal.display(dcount)

                hss_html = main_adr + 'serials/'
                dcount = count(hss_html)
                self.SeenSerialsVal.display(dcount)

                wls_html = main_adr + 'serials/wish/'
                dcount = count(wls_html)
                self.WishSerialsVal.display(dcount)

                sws_html = main_adr + 'serials/stopview/'
                dcount = count(sws_html)
                self.StoppedSerialsVal.display(dcount)

                ls_html = main_adr + 'serials/love/'
                dcount = count(ls_html)
                self.LikeSerialsVal.display(dcount)

                hsm_list = parser(hsm_html)
                wlm_list = parser(wlm_html)
                lm_list = parser(lm_html)
                hss_list = parser(hss_html)
                wls_list = parser(wls_html)
                sws_list = parser(sws_html)
                ls_list = parser(ls_html)

                directory = self.PatchToSave.text() + self.Nickname.text() + '.' + self.comboBox.currentText()

                doc = Document()

                doc.add_heading('Я посмотрел фильмы', 1)
                doc_list(hsm_list, doc)

                doc.add_heading('Я хочу посмотреть фильмы', 1)
                doc_list(wlm_list, doc)

                doc.add_heading('Любимые фильмы', 1)
                doc_list(lm_list, doc)

                doc.add_heading('Я посмотрел сериалы', 1)
                doc_list(hss_list, doc)

                doc.add_heading('Я хочу посмотреть сериалы', 1)
                doc_list(wls_list, doc)

                doc.add_heading('Я бросил смотреть сериалы', 1)
                doc_list(sws_list, doc)

                doc.add_heading('Любимые сериалы', 1)
                doc_list(ls_list, doc)

                doc.save(directory)
                QMessageBox.information(self, "Информация!", "Экспорт завершен!")

    def choose_path(self):  # Выбор пути к файлу
        dir = QFileDialog.getExistingDirectory(self, "Выбрать папку", ".")
        self.PatchToSave.setText(str(dir + '/'))
        return dir

    def get_adress_hsf(self):
        nickname = self.Nickname.text()
        mainAdr = 'https://www.movielib.ru/viewer/' + nickname + '/'
        return mainAdr



def doc_list(item_list, doc):
    rows = 1

    for i in item_list:
        rows += 1

    table = doc.add_table(rows, 2)
    table.style = 'Table Grid'
    table_items = table.rows[0].cells
    table_items[0].text = 'Название'
    table_items[1].text = 'Оригинальное название'

    row = 1

    for i in item_list:
        tablerows = table.rows[row].cells
        title = str(i['title'])
        otitle = str(i['otitle'])
        tablerows[0].text = title
        tablerows[1].text = otitle
        row += 1

def getHtml(html, params = ''):
    req = requests.get(html, headers=HEADERS, params=params)
    return req

def count(url):
    url = getHtml(url)
    soup = BeautifulSoup(url.text, 'html.parser')
    count = soup.find_all('ul', class_='filter-info')
    count = str(count)
    count = count.replace('[<ul class="filter-info">', '')
    count = count.replace(' фильмов            			        </li></ul>]', '')
    count = count.replace(' фильм            			        </li></ul>]', '')
    count = count.replace(' фильма            			        </li></ul>]', '')
    count = count.replace('<li>', '')
    count = int(count)
    return count

def pages_count(url):
    count_pages = count(url)
    pages = round(count_pages/25)
    return pages

def parser(url):
    list = []
    pc = pages_count(url)

    for i in range(1, pc + 1, 1):
        html = url + '~' + str(i)
        elements = get_info(html)
        list = list + elements

    return list

def get_info(html):
    html = getHtml(html)
    elements = []
    soup = BeautifulSoup(html.text, 'html.parser')
    items = soup.find_all('td', class_='movie')
    for item in items:
            elements.append({
                'title': item.find('a', class_='headline').get_text(),
                'otitle': item.find('span').get_text()
                })
    for element in elements:
        check = str(element)
        if check.find('   ') != -1:
            element['otitle'] = '-'
    return elements

def main():
    app = QtWidgets.QApplication(sys.argv)  # Новый экземпляр QApplication
    window = MainWindow()  # Создаём объект класса ExampleApp
    window.show()  # Показываем окно
    app.exec_()  # и запускаем приложение

if __name__ == '__main__':  # Если мы запускаем файл напрямую, а не импортируем
    main()  # то запускаем функцию main()